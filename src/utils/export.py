import os
from docx import Document
from fpdf import FPDF
import markdown2
import re

def clean_for_pdf(text: str) -> str:
    """Removes or replaces characters that fpdf2 might not handle well with default fonts."""
    # Replace common markdown bold/italic/code symbols for a cleaner text-only look in simple PDF
    text = text.replace('**', '').replace('*', '').replace('`', '')
    # Remove characters outside the Latin-1 range as FPDF's default fonts are limited
    return "".join(i if ord(i) < 256 else " " for i in text)

def export_to_docx(markdown_content: str, output_path: str):
    """Converts markdown content to a Word Document."""
    doc = Document()
    doc.add_heading('Architecture Guard Review Report', 0)
    
    # Process lines
    lines = markdown_content.split('\n')
    for line in lines:
        if line.startswith('# '):
            doc.add_heading(line[2:], level=1)
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=2)
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=3)
        elif line.startswith('- ') or line.startswith('* '):
            doc.add_paragraph(line[2:], style='List Bullet')
        elif line.strip():
            # Basic cleanup of bold/italic for docx
            clean_line = line.replace('**', '').replace('*', '').replace('`', '')
            doc.add_paragraph(clean_line)
            
    doc.save(output_path)
    return output_path

def export_to_pdf(markdown_content: str, output_path: str):
    """Converts markdown content to a PDF."""
    pdf = FPDF()
    pdf.add_page()
    pdf.set_auto_page_break(auto=True, margin=15)
    
    # Title
    pdf.set_font("Helvetica", "B", 18)
    pdf.cell(0, 15, txt="Architecture Guard Review Report", ln=True, align='C')
    pdf.ln(5)
    
    lines = markdown_content.split('\n')
    for line in lines:
        line = line.strip()
        if not line:
            pdf.ln(3)
            continue
            
        if line.startswith('# '):
            pdf.set_font("Helvetica", "B", 16)
            pdf.multi_cell(0, 10, txt=clean_for_pdf(line[2:]))
        elif line.startswith('## '):
            pdf.set_font("Helvetica", "B", 14)
            pdf.multi_cell(0, 9, txt=clean_for_pdf(line[3:]))
        elif line.startswith('### '):
            pdf.set_font("Helvetica", "B", 12)
            pdf.multi_cell(0, 8, txt=clean_for_pdf(line[4:]))
        else:
            pdf.set_font("Helvetica", size=10)
            pdf.multi_cell(0, 7, txt=clean_for_pdf(line))
            
    pdf.output(output_path)
    return output_path
